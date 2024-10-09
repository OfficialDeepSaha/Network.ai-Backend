import asyncio
import hashlib
import hmac
import json
from typing import Dict, List, Optional
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Depends, Query, Request, UploadFile, WebSocket, WebSocketDisconnect, requests , status, websockets
from fastapi.exceptions import RequestValidationError
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from flask import jsonify
from grpc import Status
import httpx
from jose import JWTError
from werkzeug.utils import escape  # Use escape if you need to quote URLs
import jwt
import numpy as np
from pydantic import BaseModel
from sqlalchemy.orm import Session , joinedload
from datetime import datetime, timedelta
import openai
from sklearn.metrics.pairwise import cosine_similarity
import uvicorn
from models import Chat, Document, GroupChat, GroupChatMember, GroupChatMessage, Message, MessageResponse, SessionLocal, User, Agent, Feedback, UserFeedback, init_db
import models
from schemas import ActionEnum, AgentModel, Connection, ConnectionRequest, ConnectionRequestResponse, ConnectionStatus, GroupChatMessageCreate, GroupChatMessageResponse, GroupChatResponse, GroupRequest, GroupRequestResponse, GroupRequestStatus, Subscription, UserCreate, UserUpdate, UserResponse, FeedbackCreate
import utils
from decision_engine import DecisionEngine
from engagement_module import EngagementModule
from retrain_model import retrain_model
from fastapi.middleware.cors import CORSMiddleware
from hashlib import sha256
import tweepy  # Added for Twitter integration
from docx import Document
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
import logging
import socketio

app = FastAPI()

# # Initialize BackgroundScheduler
scheduler = BackgroundScheduler()

logger = logging.getLogger(__name__)

# Setup CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

sio = socketio.AsyncServer(cors_allowed_origins='*')

sio_app = socketio.ASGIApp(sio, other_asgi_app=app)

openai.api_key = "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"

@sio.event
async def connect(sid, environ):
    print(f"Client connected: {sid}")

@sio.event
async def disconnect(sid):
    print(f"Client disconnected: {sid}")
# Facebook API credentials (you must fill these in with your own values)
access_token = 'EAASAZAoKzYgwBOZC8llseXB8JsSvo7GbawOUnDpAkgpgRZC0G5zIrFFdudUppT1rJCQa2mTxB8Oe0rG0W7BkiZAgQmbdZBfKHhI69Ag2ctIFPtrLUyZAbijUMPVuoX5ZBoTrRH577kuCxJIsk7ICReKDmutfv3xUKmMYGQpnlJZACioZCb3R9BapLy0H5qMiLidhet9GVSeVCYxPBIbAp9kRC7ab5qr4ZD'
app_secret = '80f37091669fd4e624b79d6cb6143672'
# Database session dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()



# Register User
from passlib.context import CryptContext


# JWT Configuration
SECRET_KEY = "8cn-qdMFKKv8q-q_fTdrr3hv9jTuphDP6zbjb5ofyjA"  # Change this to a secure random key
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 4320


# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# OAuth2 scheme
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")




# Token model
class Token(BaseModel):
    access_token: str
    token_type: str

class TokenData(BaseModel):
    email: Optional[str] = None

# User registration response model
class UserRegistrationResponse(BaseModel):
    user: UserResponse
    access_token: str
    token_type: str

# Function to verify password
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

# Function to get password hash
def get_password_hash(password):
    return pwd_context.hash(password)

# Function to authenticate user
def authenticate_user(db, email: str, password: str):
    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(password, user.password):
        return False
    return user

# Function to create access token
def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# Function to get current user
async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code= status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
        token_data = TokenData(email=email)
    except JWTError:
        raise credentials_exception
    user = db.query(User).filter(User.email == token_data.email).first()
    if user is None:
        raise credentials_exception
    return user

# New endpoint for user login
@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,

            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

# Modified endpoint to get user information using JWT
@app.get("/users/me", response_model=UserResponse)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return UserResponse.from_orm(current_user)

# Modified user registration endpoint with auto-login
@app.post("/register/", response_model=UserRegistrationResponse)
def register_user(user: UserCreate, db: Session = Depends(get_db)):
    # Check if user already exists
    db_user = db.query(User).filter(User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    hashed_password = get_password_hash(user.password)
    db_user = User(
        name=user.name,
        email=user.email,
        password=hashed_password,
        bio= user.bio
        
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)

    # Create access token for auto-login
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": db_user.email}, expires_delta=access_token_expires
    )

    # Return user info and access token
    return UserRegistrationResponse(
        user=UserResponse.from_orm(db_user),
        access_token=access_token,
        token_type="bearer"
    )






from google.oauth2 import id_token
from google.auth.transport import requests as google_requests



# Google OAuth2 Token verification
async def verify_google_token(token: str):
    try:
        id_info = id_token.verify_oauth2_token(token, google_requests.Request())
        return id_info
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid Google Token")





# Register user with Google OAuth2
# Google OAuth2 Login
@app.post("/google-login/")
async def google_login(google_token: str, db: Session = Depends(get_db)):
    id_info = await verify_google_token(google_token)
    email = id_info['email']

    db_user = db.query(User).filter(User.email == email).first()
    if not db_user:
        # Register the user if they don't exist
        db_user = User(
            name=id_info.get("name"),
            email=email,
            password="",  # Empty password, since we are using Google for authentication
        )
        db.add(db_user)
        db.commit()
        db.refresh(db_user)

    access_token = create_access_token(data={"sub": db_user.email})
    return {"access_token": access_token, "token_type": "bearer"}




import requests

GITHUB_CLIENT_ID = "Ov23liql1prInLFdkg1q"
GITHUB_CLIENT_SECRET = "59711e8713a1ad73e8a1ce42295c4d62ccca4f82"



# GitHub login request model
class GithubLoginRequest(BaseModel):
    github_token: str

# Route for GitHub login
@app.post("/auth/github")
def github_login(code: str, db: Session = Depends(get_db)):
    # Exchange authorization code for access token
    token_url = "https://github.com/login/oauth/access_token"
    headers = {"Accept": "application/json"}
    payload = {
        "client_id": GITHUB_CLIENT_ID,
        "client_secret": GITHUB_CLIENT_SECRET,
        "code": code
    }
    response = requests.post(token_url, headers=headers, data=payload)
    token_data = response.json()

    if "access_token" not in token_data:
        raise HTTPException(status_code=400, detail="GitHub OAuth failed")

    access_token = token_data["access_token"]

    # Fetch user info from GitHub API
    user_info_url = "https://api.github.com/user"
    user_info_response = requests.get(user_info_url, headers={
        "Authorization": f"Bearer {access_token}"
    })
    user_info = user_info_response.json()

    # Check if user exists in the database
    user = db.query(User).filter(User.email == user_info["email"]).first()
    
    if not user:
        # Create new user if not exists
        user = User(
            email=user_info["email"],
            name=user_info["login"],
            bio=user_info.get("bio", ""),
            password=None  # No password needed for OAuth
        )
        db.add(user)
        db.commit()
        db.refresh(user)

    # Generate JWT token
    access_token = create_access_token(data={"sub": user.email})
    return {"access_token": access_token, "token_type": "bearer" , "user_id": user.id}





# Utility function for reading .docx files
async def read_docx_file(file_path: str) -> Optional[str]:
    """
    Read the content of a DOCX file.
    
    :param file_path: Path to the DOCX file.
    :return: The content of the DOCX file as a string or None if there's an error.
    """
    try:
        logger.info(f"Reading DOCX file from path: {file_path}")
        
        # Ensure that the file is indeed a .docx file
        if not file_path.lower().endswith('.docx'):
            raise ValueError("File must be a DOCX file.")

        # Open and read the DOCX file
        document = Document(file_path)
        content = [para.text for para in document.paragraphs if para.text.strip()]
        return "\n".join(content)
    
    except Exception as e:
        logger.error(f"Error reading DOCX file: {str(e)}")
        return None

# Helper to calculate cosine similarity between multiple user embeddings
def find_common_interests(user_embeddings, threshold=0.75):
    # Calculate cosine similarity between all user embeddings
    similarity_matrix = cosine_similarity(user_embeddings)
    
    # Find pairs with similarity above the threshold
    similar_users = []
    num_users = len(user_embeddings)
    for i in range(num_users):
        for j in range(i + 1, num_users):
            if similarity_matrix[i][j] >= threshold:
                similar_users.append((i, j))
    
    return similar_users

# Optimized function for creating group chats based on common interests
# def auto_create_group_chats(similar_users, db: Session):
#     for pair in similar_users:
#         user_1 = db.query(User).get(pair[0])
#         user_2 = db.query(User).get(pair[1])

#         existing_group = db.query(GroupChat).filter(
#             GroupChat.members.any(id=user_1.id),
#             GroupChat.members.any(id=user_2.id)
#         ).first()

#         if not existing_group:
#             group_name = f"Group - {user_1.username} & {user_2.username}"
#             group_chat = GroupChat(name=group_name)
#             group_chat.members.append(user_1)
#             group_chat.members.append(user_2)
#             db.add(group_chat)
#             db.commit()

#             # Initiate AI conversation if required
#             agent_1 = db.query(Agent).filter(Agent.user_id == user_1.id).first()
#             agent_2 = db.query(Agent).filter(Agent.user_id == user_2.id).first()
#             if agent_1 and agent_2:
#                 initiate_ai_conversation(agent_1, agent_2, db)

#             group_chat_response = GroupChatResponse.from_orm(group_chat)
#             print(f"Created new group chat: {group_chat_response}")

# Run the group chat creation logic (could be scheduled or triggered after training)
def get_all_user_embeddings(db: Session):
    try:
        users = db.query(User).all()
        embeddings = [user.embedding for user in users if user.embedding is not None]

        # Debugging: Print sizes of embeddings
        sizes = [len(embedding) for embedding in embeddings]
        print(f"Embedding sizes: {sizes}")

        if embeddings:
            shape = len(embeddings[0])
            # Filter out embeddings with inconsistent sizes
            valid_embeddings = [embedding for embedding in embeddings if len(embedding) == shape]
            if len(valid_embeddings) != len(embeddings):
                print("Warning: Some embeddings had inconsistent sizes and were excluded.")

            if not valid_embeddings:
                raise ValueError("No valid embeddings found with consistent sizes")

            return valid_embeddings
        else:
            raise ValueError("No embeddings found")
    except Exception as e:
        print(f"Error fetching user embeddings: {e}")
        raise


def get_user_by_id(user_id: int, db: Session):
    return db.query(User).filter(User.id == user_id).first()

# AI-powered conversation between agents

def initiate_ai_conversation(agent_1, agent_2, db: Session):
    prompt = f"Agent {agent_1.name} and Agent {agent_2.name} are discussing their mutual interests. Start a conversation based on their shared interests."

    try:
        # Call OpenAI's GPT model to generate conversation
        conversation = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an AI conversational agent."},
                {"role": "user", "content": prompt},
            ]
        )

        response = conversation['choices'][0]['message']['content']

        # Find the group chat where the conversation should be logged
        chat = db.query(GroupChat).filter(
            GroupChat.members.any(id=agent_1.user_id),
            GroupChat.members.any(id=agent_2.user_id)
        ).first()

        if chat:
            new_message = GroupChatMessage(
                chat_id=chat.id,
                sender_id=agent_1.user_id,
                content=response,
                created_at=datetime.utcnow()
            )
            db.add(new_message)
            db.commit()

            group_chat_response = GroupChatResponse.from_orm(chat)
            return group_chat_response

    except Exception as e:
        print(f"Error initiating AI conversation: {e}")
        raise HTTPException(status_code=500, detail="Failed to initiate AI conversation")




@app.get("/get_group_chats/", response_model=List[GroupChatResponse])
def get_group_chats(user_id: int, db: Session = Depends(get_db)):
    # Query to fetch group chats where the user is a member
    group_chats = db.query(models.GroupChat).join(models.GroupChatMember).filter(
        models.GroupChatMember.user_id == user_id
    ).options(joinedload(models.GroupChat.members)).all()

    if not group_chats:
        raise HTTPException(status_code=404, detail="No group chats found for this user.")

    return [
        GroupChatResponse(
            id=group_chat.id,
            name=group_chat.name,
            created_at=group_chat.created_at,
            members=[
                UserResponse(
                    id=member.id,
                    name=member.name,
                    email=member.email  # Ensure you pass the email field here
                )
                for member in group_chat.members
            ]
        )
        for group_chat in group_chats
    ]


# Example usage when mutual interests are detected:
@app.post("/auto_group_chat/")
def trigger_group_chat_creation(db: Session = Depends(get_db)):
    try:
        user_embeddings = get_all_user_embeddings(db)
        similar_users = find_common_interests(user_embeddings)
        auto_create_group_chats(similar_users, db)
        return {"message": "Group chats and AI conversations initiated"}
    except Exception as e:
        print(f"Error triggering group chat creation: {e}")
        raise HTTPException(status_code=500, detail="Failed to create group chats and initiate conversations")


class ConnectionRequestModel(BaseModel):
    user_id: int
    target_user_id: int

# Send Connection Request to a Recommended User
@app.post("/send_connection_request/")
def send_connection_request(
    user_id: int = Query(..., description="ID of the user sending the request"),
    target_user_id: int = Query(..., description="ID of the user receiving the request"),
    db: Session = Depends(get_db)
):
    # Fetch the user's recommendations
    recommendations = get_recommendations(user_id, db)

    # Check if the target user is part of the recommended users
    if not any(rec['id'] == target_user_id for rec in recommendations):
        raise HTTPException(status_code=403, detail="You can only send connection requests to recommended users")

    # Create a new connection request
    connection_request = ConnectionRequest(
        sender_id=user_id,
        receiver_id=target_user_id,
        status=ConnectionStatus.PENDING,
        action=ActionEnum.CONNECTION_REQUEST
    )
    db.add(connection_request)
    db.commit()

    # Notify target user of the connection request
    utils.notify_user_of_connection_request(target_user_id)

    return {"status": "Connection request sent", "target_user_id": target_user_id}











@app.post("/send_connection_request_to_other/")
def send_connection_request_to_other_network(
    user_id: int = Query(..., description="ID of the user sending the request"),
    target_user_id: int = Query(..., description="ID of the user receiving the request"),
    db: Session = Depends(get_db)
):
   # Fetch the user's first-degree connections
    first_degree_connections = db.query(Connection).filter(
        (Connection.user_id_1 == user_id) | (Connection.user_id_2 == user_id)
    ).all()

    # Find second-degree connections (connections of the first-degree connections)
    second_degree_user_found = False
    first_degree_user = None
    for connection in first_degree_connections:
        first_degree_user_id = connection.user_id_1 if connection.user_id_2 == user_id else connection.user_id_2

        # Find the networks of the first-degree connection
        second_degree_connections = db.query(Connection).filter(
            (Connection.user_id_1 == first_degree_user_id) | (Connection.user_id_2 == first_degree_user_id)
        ).all()

        # Check if the target user is part of the second-degree connections
        for second_conn in second_degree_connections:
            if second_conn.user_id_1 == target_user_id or second_conn.user_id_2 == target_user_id:
                second_degree_user_found = True
                first_degree_user = first_degree_user_id
                break

    if not second_degree_user_found or not first_degree_user:
        raise HTTPException(status_code=403, detail="You can only send connection requests to second-degree users")

    # Create a new connection request for approval by the first-degree user
    approval_request = GroupRequest(
        sender_id=user_id,
        receiver_id=first_degree_user_id,  # This goes to the first-degree user for approval
        target_user_id=target_user_id,  # The second-degree target user is stored
        status=GroupRequestStatus.PENDING,
        action=ActionEnum.CONNECTION_REQUEST
    )
    db.add(approval_request)
    db.commit()

    # Notify the first-degree user of the approval request
    utils.notify_user_of_connection_request(target_user_id)

    return {"status": "Connection request sent for approval", "target_user_id": target_user_id, "first_degree_user": first_degree_user}









def get_current_user_id(user_id: int = Query(...)) -> int:
    return user_id

@app.get("/approval_requests/", response_model=List[ConnectionRequestResponse])
def get_approval_requests(user_id: int = Query(...), db: Session = Depends(get_db)):
    # Fetch connection requests for the given user
    connection_requests = db.query(ConnectionRequest).filter(
        ConnectionRequest.receiver_id == user_id,
        ConnectionRequest.status == ConnectionStatus.PENDING
    ).all()
    
    if not connection_requests:
        return []

    response = []
    for request in connection_requests:
        # Fetch the sender's name
        sender = db.query(User).filter(User.id == request.sender_id).first()  # Assuming User is the model for user data

        if sender is None:
            sender_name = "Unknown"  # Handle the case where sender does not exist
        else:
            sender_name = sender.name  # Assuming 'name' is the column for the sender's name in User model
        
        response.append(ConnectionRequestResponse(
            id=request.id,
            sender_id=request.sender_id,
            sender_name=sender_name,  # Include the sender's name in the response
            receiver_id=request.receiver_id,
            status=request.status,
            created_at=request.created_at,
            updated_at=request.updated_at
        ))

    return response






@app.get("/approval_requests_for_group/", response_model=List[GroupRequestResponse])
def get_approval_requests_for_group_creation(user_id: int = Query(...), db: Session = Depends(get_db)):
    # Fetch group requests for the user
    connection_requests = db.query(GroupRequest).filter(
        GroupRequest.receiver_id == user_id,
        GroupRequest.status == GroupRequestStatus.PENDING  # Only fetch pending requests
    ).all()

    if not connection_requests:
        return []

    response = []
    for request in connection_requests:
        # Find the sender's name using the sender_id
        sender = db.query(User).filter(User.id == request.sender_id).first()
        sender_name = sender.name if sender else "Unknown"

        # Append to the response
        response.append(GroupRequestResponse(
            id=request.id,
            sender_id=request.sender_id,
            sender_name=sender_name,  # Use the fetched sender name
            receiver_id=request.receiver_id,
            target_user_id=request.target_user_id,
            status=request.status.value,  # Convert enum to string
            created_at=request.created_at,
            updated_at=request.updated_at
        ))

    return response




@app.post("/auto_generate_connections/")
def auto_generate_connections(user_id: int, db: Session = Depends(get_db)):
    # Fetch recommendations using the embeddings-based recommendation system
    recommendations = get_recommendations(user_id, db)
    
    for rec in recommendations:
        target_user_id = rec['id']  # This is the recommended user
        similarity_score = rec['similarity']  # Fetch similarity score

        # Only send connection requests if similarity score is greater than 0.9
        if similarity_score > 0.95:
            # Check if a connection request already exists from the recommended user to the main user
            existing_request = db.query(ConnectionRequest).filter(
                ConnectionRequest.sender_id == target_user_id,  # Reversed sender
                ConnectionRequest.receiver_id == user_id,  # Reversed receiver
                ConnectionRequest.status == ConnectionStatus.PENDING
            ).first()

            if not existing_request:
                # The recommended user (target_user_id) sends the connection request to the main user (user_id)
                connection_request = ConnectionRequest(
                    sender_id=target_user_id,  # Reversed sender
                    receiver_id=user_id,  # Reversed receiver
                    status=ConnectionStatus.PENDING,
                    action=ActionEnum.CONNECTION_REQUEST
                )
                db.add(connection_request)
                db.commit()

                # Notify the main user that a connection request has been received from the recommended user
                utils.notify_user_of_connection_request(user_id)

    return {"status": "Auto-generated connection requests for similar users sent to main user"}





# Handle approval actions for connection requests
import logging

logger = logging.getLogger(__name__)

@app.post("/handle_approval/{approval_request_id}", response_model=ConnectionRequestResponse)
def handle_approval(
    approval_request_id: int,
    approved: bool = Query(..., description="Approval status"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # Get current_user from JWT token
):
    # Query the approval request from the database
    approval_request = db.query(ConnectionRequest).filter(ConnectionRequest.id == approval_request_id).first()
    sender = db.query(User).filter(User.id == approval_request.sender_id).first()
    sender_name = sender.name if sender else "Unknown"
    if not approval_request:
        raise HTTPException(status_code=404, detail="Approval request not found")
    


    # Check if the current user is authorized to handle this approval
    if approval_request.receiver_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized to handle this request")

    if approved:
        approval_request.status = ConnectionStatus.ACCEPTED



        # Add the connection to the user's network
        new_connection = Connection(
            user_id_1=approval_request.sender_id,
            user_id_2=approval_request.receiver_id,
            created_at=datetime.utcnow()
        )
        db.add(new_connection)

        # Automatically create a chat if approved
        existing_chat = db.query(Chat).filter(
            (Chat.user_1_id == approval_request.sender_id) & (Chat.user_2_id == approval_request.receiver_id) |
            (Chat.user_1_id == approval_request.receiver_id) & (Chat.user_2_id == approval_request.sender_id)
        ).first()

        if not existing_chat:
            chat = Chat(user_1_id=approval_request.sender_id, user_2_id=approval_request.receiver_id)
            db.add(chat)

        # # Additional logic for group chats and agent conversation
        # if approval_request.status == ConnectionStatus.ACCEPTED:
        #     auto_create_group_chats([(approval_request.sender_id, approval_request.receiver_id)], db)
            

        elif approval_request.action == ActionEnum.START_CONVERSATION:
            agent_1 = db.query(Agent).filter(Agent.user_id == approval_request.sender_id).first()
            agent_2 = db.query(Agent).filter(Agent.user_id == approval_request.receiver_id).first()
            if agent_1 and agent_2:
                initiate_ai_conversation(agent_1, agent_2, db)
            else:
                raise HTTPException(status_code=404, detail="Agents not found for conversation initiation")
    else:
        approval_request.status = ConnectionStatus.REJECTED

    approval_request.updated_at = datetime.utcnow()
    db.commit()

    return ConnectionRequestResponse(
        id=approval_request.id,
        sender_id=approval_request.sender_id,
        sender_name=sender_name,
        receiver_id=approval_request.receiver_id,
        status=approval_request.status,
        created_at=approval_request.created_at,
        updated_at=approval_request.updated_at
    )







@app.post("/handle_approval_other_network/{approval_request_id}", response_model=GroupRequestResponse)
def handle_approval_other_network(
    approval_request_id: int,
    approved: bool = Query(..., description="Approval status"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)  # Get current_user from JWT token
):
    # Query the approval request from the database
    approval_request = db.query(GroupRequest).filter(GroupRequest.id == approval_request_id).first()
    if not approval_request:
        raise HTTPException(status_code=404, detail="Approval request not found")

    # Check if the current user is authorized to handle this approval
    if approval_request.receiver_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized to handle this request")

    if approved:
        approval_request.status = GroupRequestStatus.ACCEPTED

        # Add the connection between the original sender (user_id) and the second-degree user (target_user_id)
        new_connection = Connection(
            user_id_1=approval_request.sender_id,   # The original user who sent the request
            user_id_2=approval_request.target_user_id,  # The second-degree user who is the target
            created_at=datetime.utcnow()
        )
        db.add(new_connection)

        # Automatically create a chat if approved
        existing_chat = db.query(Chat).filter(
            (Chat.user_1_id == approval_request.sender_id) & (Chat.user_2_id == approval_request.target_user_id) |
            (Chat.user_1_id == approval_request.target_user_id) & (Chat.user_2_id == approval_request.sender_id)
        ).first()

        if not existing_chat:
            chat = Chat(user_1_id=approval_request.sender_id, user_2_id=approval_request.target_user_id)
            db.add(chat)

        # Create a group with the current user (first-degree), the second-degree target, and the original sender
        auto_create_group_chats(approval_request.sender_id, current_user.id, approval_request.target_user_id, db)

        # Notify sender (user who sent the request)
        sender = db.query(User).filter(User.id == approval_request.sender_id).first()
        sender_name = sender.name if sender else "Unknown"

    else:
        approval_request.status = GroupRequestStatus.REJECTED

    approval_request.updated_at = datetime.utcnow()
    db.commit()

    return GroupRequestResponse(
        id=approval_request.id,
        sender_id=approval_request.sender_id,
        receiver_id=approval_request.receiver_id,
        sender_name=sender_name,
        target_user_id=approval_request.target_user_id,
        status=approval_request.status.value,
        created_at=approval_request.created_at,
        updated_at=approval_request.updated_at
    )
















@app.get("/network/{user_id}", response_model=List[UserResponse])
def get_user_network(user_id: int, db: Session = Depends(get_db)):
    # Fetch connections for the user
    connections = db.query(Connection).filter(
        (Connection.user_id_1 == user_id) | (Connection.user_id_2 == user_id)
    ).all()

    # Prepare the response with connected users' info
    user_network = []
    for connection in connections:
        # Identify the connected user
        connected_user_id = connection.user_id_2 if connection.user_id_1 == user_id else connection.user_id_1
        connected_user = db.query(User).filter(User.id == connected_user_id).first()
        
        if connected_user:
            # Append to user_network ensuring `id` is passed, not `user_id`
            user_network.append(UserResponse(
                id=connected_user.id,  # Correct field name
                name=connected_user.name,
                email=connected_user.email,
                bio=connected_user.bio,  # Ensure bio is fetched and passed
                education=connected_user.education,
                experience=connected_user.experience,
                goal=connected_user.goal,
                github=connected_user.github
            ))

    return user_network







# Define a new response model for a clearer response structure
class NetworkOfNetworkResponse(BaseModel):
    first_degree_user: UserResponse  # The direct connection whose network you're viewing
    second_degree_users: List[UserResponse]  # List of their second-degree connections


@app.get("/network_of_networks/{user_id}", response_model=List[NetworkOfNetworkResponse])
def get_network_of_networks(user_id: int, db: Session = Depends(get_db)):
    """
    Fetch the networks of the user's direct connections (i.e., second-degree connections)
    and indicate whose network each second-degree connection belongs to.
    """
    # Step 1: Fetch the user's direct connections (first-degree network)
    connections = db.query(Connection).filter(
        (Connection.user_id_1 == user_id) | (Connection.user_id_2 == user_id)
    ).all()

    # Prepare a set to avoid duplicates in first-degree connections
    first_degree_user_ids = set()

    # Collect all first-degree connection user IDs
    for connection in connections:
        connected_user_id = connection.user_id_2 if connection.user_id_1 == user_id else connection.user_id_1
        first_degree_user_ids.add(connected_user_id)

    # Step 2: Fetch second-degree connections for each first-degree user
    network_of_networks = []

    for connection_user_id in first_degree_user_ids:
        # Fetch the first-degree user's details
        first_degree_user = db.query(User).filter(User.id == connection_user_id).first()
        
        # Fetch the connections of the connected users (second-degree)
        second_degree_connections = db.query(Connection).filter(
            (Connection.user_id_1 == connection_user_id) | (Connection.user_id_2 == connection_user_id)
        ).all()

        # Prepare the second-degree users for this first-degree connection
        second_degree_users = []
        for second_connection in second_degree_connections:
            second_degree_user_id = second_connection.user_id_2 if second_connection.user_id_1 == connection_user_id else second_connection.user_id_1
            # Exclude the original user and their direct connections
            if second_degree_user_id != user_id and second_degree_user_id not in first_degree_user_ids:
                second_degree_user = db.query(User).filter(User.id == second_degree_user_id).first()
                if second_degree_user:
                    second_degree_users.append(UserResponse(
                        id=second_degree_user.id,
                        name=second_degree_user.name,
                        email=second_degree_user.email,
                        bio=second_degree_user.bio or None,
                        education=second_degree_user.education or None,
                        experience=second_degree_user.experience or None,
                        goal=second_degree_user.goal or None,
                        github=second_degree_user.github or None
                    ))

        # Append the first-degree user and their second-degree users to the response
        network_of_networks.append(NetworkOfNetworkResponse(
            first_degree_user=UserResponse(
                id=first_degree_user.id ,
                name=first_degree_user.name,
                email=first_degree_user.email,
                bio=first_degree_user.bio or None,
                education=first_degree_user.education or None,
                experience=first_degree_user.experience or None,
                goal=first_degree_user.goal or None,
                github=first_degree_user.github or None
            ),
            second_degree_users=second_degree_users
        ))

    return network_of_networks










class MessageRequest(BaseModel):
    chat_id: int
    sender_id: int
    content: str




@app.post("/send_message/", response_model=MessageResponse)
async def send_message(
    message_request: MessageRequest,  # Use the Pydantic model for request body
    db: Session = Depends(get_db)
):
    chat = db.query(Chat).filter(Chat.id == message_request.chat_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")

    # Create a new message
    message = Message(
        chat_id=message_request.chat_id,
        sender_id=message_request.sender_id,
        content=message_request.content
    )
    db.add(message)
    db.commit()

    # Broadcast the message to all users in the chat via WebSocket (to be implemented)
    await broadcast_message_to_chat(message.chat_id, message.content)

    return MessageResponse(
        id=message.id,
        chat_id=message.chat_id,
        sender_id=message.sender_id,
        content=message.content,
        created_at=message.created_at
    )

@app.get("/get_messages/{chat_id}", response_model=List[MessageResponse])
def get_messages(chat_id: int, db: Session = Depends(get_db)):
    chat = db.query(Chat).filter(Chat.id == chat_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")

    messages = db.query(Message).filter(Message.chat_id == chat_id).order_by(Message.created_at.asc()).all()

    return [
        MessageResponse(
            id=message.id,
            chat_id=message.chat_id,
            sender_id=message.sender_id,
            content=message.content,
            created_at=message.created_at
        ) for message in messages
    ]



@app.get("/get_chat_id")
def get_chat_id(user_id: int, db: Session = Depends(get_db)):
    # Fetch all chats involving the user
    chats = db.query(Chat).filter(
        (Chat.user_1_id == user_id) | (Chat.user_2_id == user_id)
    ).all()

    if not chats:
        raise HTTPException(status_code=404, detail="No chats found for the given user ID")

    # Prepare the response to hold multiple chats
    chat_details = []

    # Iterate over each chat and fetch user details
    for chat in chats:
        # Fetch the names of user_1 and user_2 from the User table
        user_1 = db.query(User).filter(User.id == chat.user_1_id).first()
        user_2 = db.query(User).filter(User.id == chat.user_2_id).first()

        if not user_1 or not user_2:
            raise HTTPException(status_code=404, detail="One or both users not found for a chat")

        # Append each chat's details to the chat_details list
        chat_details.append({
            "chat_id": chat.id,
            "user_1_id": chat.user_1_id,
            "user_1_name": user_1.name,
            "user_2_id": chat.user_2_id,
            "user_2_name": user_2.name,
            "created_at": chat.created_at
        })

    # Return the list of chats with the relevant details
    return chat_details



active_connections = {}


@app.websocket("/ws/{chat_id}")
async def websocket_endpoint(websocket: WebSocket, chat_id: int, token: str = Query(...), db: Session = Depends(get_db)):
    # Manually validate the token
    user_id = await get_current_user(token , db)
    
    logging.info(f"User {user_id} connected to chat {chat_id}")

    await websocket.accept()

    # Register the user in active connections
    if chat_id not in active_connections:
        active_connections[chat_id] = []
    active_connections[chat_id].append(websocket)

    try:
        while True:
            data = await websocket.receive_text()
            # Broadcast received message to all users in the chat
            await broadcast_message_to_chat(chat_id, data)
    except WebSocketDisconnect:
        # Remove the WebSocket connection on disconnect
        active_connections[chat_id].remove(websocket)



async def broadcast_message_to_chat(chat_id: int, message: str):
    if chat_id in active_connections:
        for connection in active_connections[chat_id]:
            try:
                await connection.send_text(message)
            except Exception as e:
                # Handle any WebSocket errors here
                logging.error(f"Error sending message via WebSocket: {str(e)}")






@app.get("/my_groups/", response_model=List[GroupChatResponse])
def get_my_groups(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    group_chats = db.query(GroupChat).join(GroupChatMember).filter(GroupChatMember.user_id == current_user.id).all()

    group_chat_data = []
    for group_chat in group_chats:
        latest_message = db.query(GroupChatMessage).filter(
            GroupChatMessage.chat_id == group_chat.id
        ).order_by(GroupChatMessage.created_at.desc()).first()

        # Query members with additional fields like email and bio
        members = db.query(User).join(GroupChatMember).filter(
            GroupChatMember.group_chat_id == group_chat.id
        ).all()

        # Ensure members have their email and bio fetched as well
        group_chat_data.append({
            "id": group_chat.id,
            "name": group_chat.name,
            "created_at": group_chat.created_at,
            "members": [
                UserResponse(id=member.id, name=member.name, email=member.email, bio=member.bio) for member in members
            ],
            "latest_message": latest_message.content if latest_message else None
        })

    if not group_chat_data:
        raise HTTPException(status_code=404, detail="No groups found for the current user")

    return group_chat_data







from sqlalchemy.exc import SQLAlchemyError

# Track active connections by chat_id
activate_connections = {}

@app.websocket("/ws/{chat_id}")
async def websocket_endpoint_for_groups(websocket: WebSocket, chat_id: int, token: str = Query(...), db: Session = Depends(get_db)):
    # Manually validate the token and get the user
    user_id = await get_current_user(token, db)
    
    logging.info(f"User {user_id} connected to chat {chat_id}")

    await websocket.accept()

    # Register the user in active connections
    if chat_id not in activate_connections:
        activate_connections[chat_id] = []
    activate_connections[chat_id].append(websocket)

    try:
        while True:
            data = await websocket.receive_json()

            # Extract message content from data
            message_content = data.get("content")

            # Save the message to the database
            try:
                message = GroupChatMessage(
                    chat_id=chat_id,
                    sender_id=user_id,
                    content=message_content,
                    created_at=datetime.utcnow()
                )
                db.add(message)

                # Commit the transaction to save the message
                logging.info(f"Attempting to save message to chat {chat_id} from user {user_id}")
                db.commit()
                db.refresh(message)
                logging.info(f"Message saved to chat {chat_id} with message ID {message.id}")

            except SQLAlchemyError as e:
                db.rollback()  # Rollback the session if there’s an error
                logging.error(f"Error saving message to database: {str(e)}")
                continue  # Don't break the loop, continue handling new messages

            # Broadcast received message to all users in the chat
            await broadcast_message_to_group(chat_id, message)

    except WebSocketDisconnect:
        # Remove the WebSocket connection on disconnect
        activate_connections[chat_id].remove(websocket)
        logging.info(f"User {user_id} disconnected from chat {chat_id}")


async def broadcast_message_to_group(chat_id: int, message: GroupChatMessage):
    if chat_id in activate_connections:
        # Prepare message response
        message_data = {
            "id": message.id,
            "chat_id": message.chat_id,
            "sender_id": message.sender_id,
            "content": message.content,
            "created_at": message.created_at.isoformat()
        }

        # Broadcast message to all connected clients in this chat
        for connection in activate_connections[chat_id]:
            try:
                await connection.send_json(message_data)
            except Exception as e:
                logging.error(f"Error sending message via WebSocket: {str(e)}")



@app.get("/group_chats/{chat_id}/messages", response_model=List[GroupChatMessageResponse])
def get_chat_messages(chat_id: int, db: Session = Depends(get_db)):
    messages = db.query(GroupChatMessage).filter(GroupChatMessage.chat_id == chat_id).order_by(GroupChatMessage.created_at.asc()).all()

    if not messages:
        raise HTTPException(status_code=404, detail="No messages found for this chat")

    return messages



@app.post("/api/askAI")
async def ask_ai(request: Request):
    data = await request.json()
    user_message = data["message"]
    chat_partner_name = data["chat_partner_name"]  # Add chat partner's name

    try:
        # Customize the system prompt to make the AI behave like a chat participant
        response = openai.ChatCompletion.create(
            model="gpt-4",  # or "gpt-3.5-turbo"
            messages=[
                {"role": "system", "content": f"You are {chat_partner_name}. Respond to the user like you are a real person in the chat."},
                {"role": "user", "content": user_message}  # The user's message
            ],
            max_tokens=100,
            temperature=0.8,
        )

        ai_response = response["choices"][0]["message"]["content"].strip()
        return {"response": ai_response}

    except openai.error.InvalidRequestError as e:
        return {"error": str(e)}






def auto_create_group_chats(sender_id: int, approver_id: int, target_user_id: int, db: Session):
 
# Fetch user details

    sender = db.query(User).filter(User.id == sender_id).first()
    approver = db.query(User).filter(User.id == approver_id).first()
    target_user = db.query(User).filter(User.id == target_user_id).first()

    if not sender or not approver or not target_user:
        raise HTTPException(status_code=404, detail="One or more users not found")

    # Create the group name
    group_name = f"Group - {sender.name}, {approver.name}, and {target_user.name}"

    # Create a new group chat
    group_chat = GroupChat(name=group_name)

    # Add all three members (sender, first-degree user, and second-degree user)
    group_chat.members.append(sender)
    group_chat.members.append(approver)
    group_chat.members.append(target_user)

    db.add(group_chat)
    db.commit()

            # # AI agents initiate conversation in the group chat
            # agent_1 = db.query(Agent).filter(Agent.user_id == user_1.id).first()
            # agent_2 = db.query(Agent).filter(Agent.user_id == user_2.id).first()
            # if agent_1 and agent_2:
            #     initiate_ai_conversation(agent_1, agent_2, db)


# Fetch tweets from user's Twitter handle
def fetch_twitter_data(user_handle: str):
    auth = tweepy.OAuth1UserHandler(consumer_key, consumer_secret, access_token, access_token_secret)
    api = tweepy.API(auth)
    tweets = api.user_timeline(screen_name=user_handle, count=100, tweet_mode="extended")
    tweet_texts = [tweet.full_text for tweet in tweets]
    return tweet_texts


# Aggregation function remains the same but with additional error handling
def aggregate_embeddings(embedding_list):
    if embedding_list:
        try:
            embeddings_array = np.array(embedding_list)
            aggregated_embedding = np.mean(embeddings_array, axis=0)
            return aggregated_embedding.tolist()  # Ensure this returns a list
        except Exception as e:
            logger.error(f"Error aggregating embeddings: {str(e)}")
            return None
    return None


GITHUB_BASE_URL = "https://api.github.com"
GITHUB_TOKEN = "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"



# Function to fetch GitHub user ID based on username
async def get_github_user_id(username: str) -> Optional[int]:
    """
    Fetch GitHub user ID using the username.
    
    :param username: The GitHub username.
    :return: The GitHub user ID or None if not found.
    """
    url = f"{GITHUB_BASE_URL}/users/{username}"
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
            return data.get('id') if 'id' in data else None
        except httpx.HTTPStatusError as e:
            print(f"HTTP error fetching GitHub user ID: {e}")
            return None
        except Exception as e:
            print(f"Error fetching GitHub user ID: {e}")
            return None

# Function to fetch user repositories
async def fetch_github_repos(username: str) -> Optional[List[Dict]]:
    """
    Fetch the GitHub user's public repositories.
    
    :param username: GitHub username.
    :return: List of user's public repositories or None if request fails.
    """
    url = f"{GITHUB_BASE_URL}/users/{username}/repos"
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
        
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            return response.json()
        except httpx.HTTPStatusError as e:
            print(f"HTTP error fetching GitHub repos: {e}")
            return None
        except Exception as e:
            print(f"Error fetching GitHub repos: {e}")
            return None

# Function to fetch user events/activity
async def fetch_github_user_activity(username: str) -> Optional[List[Dict]]:
    """
    Fetch recent activities of the GitHub user.
    
    :param username: GitHub username.
    :return: List of recent activities or None if request fails.
    """
    url = f"{GITHUB_BASE_URL}/users/{username}/events"
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            return response.json()
        except httpx.HTTPStatusError as e:
            print(f"HTTP error fetching GitHub activity: {e}")
            return None
        except Exception as e:
            print(f"Error fetching GitHub activity: {e}")
            return None

# Auto-train agent with GitHub data and document
async def auto_train_agent(db: Session, agent: models.Agent, github: str, doc_file: str):
    logger.info(f"Running auto-training for agent {agent.user_id}")

    try:
        # Fetch GitHub data and read document file
        github_repos = await fetch_github_repos(github)
        github_activity = await fetch_github_user_activity(github)
        if not github_repos or not github_activity:
            logger.error(f"Failed to fetch GitHub data for {github}")
            return "Failed to fetch GitHub data."

        # Ensure all repo names and activity types are strings
        combined_text = "\n".join([str(repo['name']) for repo in github_repos] + 
                                  [str(event['type']) for event in github_activity])
        # print("This is combined text:- " + combined_text)

        github_embeddings = utils.generate_embeddings(combined_text)
        # Convert the embeddings to a string for printing
        # print("This is Github Embeddings:- " + str(github_embeddings))

        # Read and generate embeddings from document
        doc_content = await read_docx_file(doc_file)
        if not doc_content:
            logger.error(f"Failed to read document {doc_file}")
            return "Failed to read document."

        doc_embeddings = utils.generate_embeddings(doc_content)
        # Convert the embeddings to a string for printing
        # print("This is Doc Embeddings- " + str(doc_embeddings))

        # Aggregate embeddings and update the agent's knowledge
        combined_embeddings = aggregate_embeddings([github_embeddings, doc_embeddings])
        # Convert the combined embeddings to a string for printing
        # print("This is Combined Embeddings:- " + str(combined_embeddings))

        if combined_embeddings:


            agentX = db.query(models.Agent).filter(models.Agent.user_id == agent.user_id).first()
            if not agent:
                logger.error(f"Agent with user_id {agent.user_id} not found.")
                return "Agent not found."

            # Ensure combined_embeddings is a list or a serializable format
            if isinstance(combined_embeddings, list):

                agentX.update_knowledge = json.dumps(combined_embeddings)  # Convert to JSON string
            else:
                logger.error(f"Unexpected format for combined_embeddings: {type(combined_embeddings)}")
                return "Unexpected format for combined embeddings."
            
            db.commit()  # Commit changes to the database
            logger.info(f"Agent {agent.user_id} auto-trained successfully.")
            return "Agent auto-trained with GitHub data and document."
        else:
            logger.error(f"Failed to aggregate embeddings for agent {agent.user_id}")
            return "Failed to aggregate embeddings."

    except Exception as e:
        logger.error(f"Error during auto-training: {str(e)}")
        return "Auto-training failed due to an error."



# Improved scheduling function
def schedule_auto_training(user_id: int, github_username: str, doc_file: str, db: Session):
    # Fetch or create the agent
    agent = db.query(Agent).filter(Agent.user_id == user_id).first()
    if not agent:
        agent = Agent(
            user_id=user_id,
            state='initial',
            last_active=datetime.now(),
            next_action=None
        )
        db.add(agent)
        db.commit()
    else:
        agent.last_active = datetime.now()
        db.commit()

    logger.info(f"Scheduling auto-training for agent {user_id}")

    try:
        # Ensure the arguments are passed in the right format to the job
        scheduler.add_job(
            auto_train_agent,
            IntervalTrigger(hours=24),
            args=[db, agent, github_username, doc_file],  # Corrected arguments
            id=f"auto_train_agent_{user_id}",
            replace_existing=True
        )
        
    except Exception as e:
        logger.error(f"Error scheduling auto-training: {str(e)}")





































# Login User
class LoginRequest(BaseModel):
    email: str
    password: str

@app.post("/login/")
def login_user(request: LoginRequest, db: Session = Depends(get_db)):
    hashed_password = sha256(request.password.encode()).hexdigest()
    db_user = db.query(User).filter(User.email == request.email, User.password == hashed_password).first()
    if db_user:
        return {"message": "Login successful", "id": db_user.id}
    else:
        raise HTTPException(status_code=401, detail="Invalid credentials")


# Improved user update function with better error handling and logging
# Configure logging
logger = logging.getLogger("uvicorn.error")

@app.post("/update_user/")
async def update_user(
    user: UserUpdate, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    db_user = db.query(models.User).filter(models.User.id == user.user_id).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")

    try:
        # Update user details if provided
        if user.education is not None:
            db_user.education = user.education
        if user.experience is not None:
            db_user.experience = user.experience
        if user.goal is not None:
            db_user.goal = user.goal
        if user.twitter_handle is not None:
            db_user.twitter_handle = user.twitter_handle
        if user.about is not None:
            db_user.about = user.about
        if user.bio is not None:
            db_user.bio = user.bio
        if user.github is not None:
            db_user.github = user.github
        if user.linkedin is not None:
            db_user.linkedin= user.linkedin
        if user.banner_image is not None:
            db_user.banner_image= user.banner_image
        if user.profile_image is not None:
            db_user.profile_image= user.profile_image                        


        # Generate embeddings for the user's profile
        user_profile_text = (
            f"Education: {db_user.education}\n"
            f"Experience: {db_user.experience}\n"
            f"Goals: {db_user.goal}\n"
            f"github: {db_user.github}"
        )

        # Generate embeddings synchronously
        user_embedding = utils.generate_embeddings(user_profile_text)
        if user_embedding:
            db_user.embedding = user_embedding
            db.commit()
        else:
            logger.error(f"Failed to generate embeddings for user {user.user_id}")
            raise HTTPException(status_code=500, detail="Error generating embeddings")

        # Ensure agent is created or updated
        agent = db.query(models.Agent).filter(models.Agent.user_id == user.user_id).first()
        if not agent:
            agent = models.Agent(user_id=user.user_id)
            db.add(agent)
            db.commit()

        # Get the latest document file path for the user (if applicable)
        document = db.query(models.Document).filter(models.Document.user_id == user.user_id).order_by(models.Document.id.desc()).first()
        if document:
            doc_file_path = document.file_path
        else:
            # Handle the case where no document is found
            doc_file_path = "path/to/default/document.docx"  # Adjust as needed

        # Schedule auto-training
        background_tasks.add_task(
            schedule_auto_training,
            user.user_id,
            user.github,
            doc_file_path,
            db
        )

        return {"message": "User details and embeddings updated successfully"}

    except Exception as e:
        logger.error(f"Error updating user: {str(e)}")
        raise HTTPException(status_code=500, detail="Error updating user details.")
  


# Collect Feedback
@app.post("/feedback/")
def collect_feedback(feedback: FeedbackCreate, db: Session = Depends(get_db)):
    db_feedback = Feedback(user_id=feedback.user_id, feedback_text=feedback.text, rating=feedback.rating)
    db.add(db_feedback)
    db.commit()
    return {"message": "Feedback received"}

# Get Recommendations (Only recommended users are eligible for connection requests)
@app.get("/recommendations/")
def get_recommendations(user_id: int, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Fetch all connection requests related to the user
    connections_requests = db.query(ConnectionRequest).filter(
        (ConnectionRequest.sender_id == user_id) | 
        (ConnectionRequest.receiver_id == user_id)
    ).filter(
        ConnectionRequest.status.in_([ConnectionStatus.ACCEPTED, ConnectionStatus.REJECTED])
    ).all()

    # Create a set of user IDs who have approved or rejected connection requests
    connected_user_ids = set()
    for conn_req in connections_requests:
        if conn_req.sender_id == user_id:
            connected_user_ids.add(conn_req.receiver_id)
        else:
            connected_user_ids.add(conn_req.sender_id)

    # Fetch all users that are directly connected via 'Connection' table
    connections = db.query(Connection).filter(
        (Connection.user_id_1 == user_id) | 
        (Connection.user_id_2 == user_id)
    ).all()

    # Add users from the 'connections' table to the connected_user_ids set
    for conn in connections:
        if conn.user_id_1 == user_id:
            connected_user_ids.add(conn.user_id_2)
        else:
            connected_user_ids.add(conn.user_id_1)

    # Include all users except those in the connected_user_ids set and the current user
    other_users = db.query(User).filter(User.id != user_id).filter(
        User.id.notin_(connected_user_ids)
    ).all()

    # Generate embeddings for the current user's profile
    user_profile_text = f"Education: {user.education}\nExperience: {user.experience}\nGoals: {user.goal}"
    user_embedding = utils.generate_embeddings(user_profile_text)

    # Generate embeddings for other users' profiles
    other_user_profiles = [
        f"Education: {other_user.education}\nExperience: {other_user.experience}\nGoals: {other_user.goal}"
        for other_user in other_users
    ]

    other_user_embeddings = utils.batch_generate_embeddings(other_user_profiles)
    recommendations = []

    # Calculate similarity between the current user and others
    for other_user, other_user_embedding in zip(other_users, other_user_embeddings):
        similarity_score = utils.calculate_similarity(user_embedding, other_user_embedding)
        if similarity_score > 0.87:  # Similarity threshold
            recommendations.append({
                "id": other_user.id,
                "name": other_user.name,
                "similarity": similarity_score,
                "bio": other_user.bio,
                "about" : other_user.about,
                "github": other_user.github,
                "linkedin": other_user.linkedin,
                "banner_image" : other_user.banner_image,
                "profile_image" : other_user.profile_image,
                "twitter_handle" : other_user.twitter_handle
                
                
            })

    # Sort recommendations by similarity in descending order
    recommendations.sort(key=lambda x: x['similarity'], reverse=True)
    return recommendations












@app.post("/upload_audio/")
async def process_audio(
    user_id: int,
    audio: UploadFile = File(...),
    db: Session = Depends(get_db)  # Adjust this import based on your project structure
):
    print("Audio processing started...")

    # Step 1: Read audio file bytes
    audio_bytes = await audio.read()  # Ensure you use the correct variable here

    # Save the audio file temporarily for the OpenAI API
    temp_audio_path = "temp_audio.mp3"
    with open(temp_audio_path, "wb") as temp_file:
        temp_file.write(audio_bytes)

    # Step 2: Transcribe Audio using Whisper
    try:
        with open(temp_audio_path, "rb") as audio_file:
            response = openai.Audio.transcribe(
                model="whisper-1",  # Use the appropriate model as required
                file=audio_file
            )
        transcription_text = response['text']  # Ensure the key access is correct
    except openai.APIError as e:
        raise HTTPException(status_code=500, detail=f"OpenAI API error: {e}")

    # Step 3: Use GPT-4 to extract profile details from transcription
    prompt = f"""
    The following text is a user's spoken description of their profile. 
    Extract their name, age, education, experience, and goals from the text, regardless of the format used. 
    
    Text: "{transcription_text}"
    
    Return the result in a JSON format with the keys: name, age, education, experience, goal.
    If any information is missing or unclear, return null for that field.
    """
    
    try:
        gpt_response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=300
        )
        
        extracted_data = gpt_response['choices'][0]['message']['content'].strip()
        
        if not extracted_data:
            return {"message": "Could not extract profile data from the transcription"}
        
        # Safely parse extracted data
        extracted_data = json.loads(extracted_data)

    except openai.APIError as e:
        raise HTTPException(status_code=500, detail=f"OpenAI GPT-4 error: {e}")
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Failed to decode extracted data")

    # Step 4: Update User Profile in Database
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Update user's profile details with extracted information
    user.name = extracted_data.get("name", user.name)
    user.education = extracted_data.get("education", user.education)
    user.experience = extracted_data.get("experience", user.experience)
    user.goal = extracted_data.get("goal", user.goal)

    db.commit()  # Commit the changes to the database

    # Step 5: Generate New Embeddings for User
    user_profile_text = f"Education: {user.education}\nExperience: {user.experience}\nGoals: {user.goal}"
    user_embedding = utils.generate_embeddings(user_profile_text)

    # Step 6: Fetch Other Users and Generate Embeddings
    other_users = db.query(User).filter(User.id != user_id).all()
    other_user_profiles = [
        f"Education: {other_user.education}\nExperience: {other_user.experience}\nGoals: {other_user.goal}"
        for other_user in other_users
    ]
    other_user_embeddings = utils.batch_generate_embeddings(other_user_profiles)

    # Step 7: Calculate Similarities and Recommend Users
    recommendations = []
    for other_user, other_user_embedding in zip(other_users, other_user_embeddings):
        similarity_score = utils.calculate_similarity(user_embedding, other_user_embedding)
        if similarity_score > 0.87:  # Similarity threshold
            recommendations.append({
                "id": other_user.id,
                "name": other_user.name,
                "similarity": similarity_score,
                "bio": other_user.bio,
                "about": other_user.about,
                "github": other_user.github,
                "linkedin": other_user.linkedin,
                "banner_image": other_user.banner_image,
                "profile_image": other_user.profile_image,
                "twitter_handle": other_user.twitter_handle
            })

    # Sort recommendations by similarity in descending order
    recommendations.sort(key=lambda x: x['similarity'], reverse=True)
    
    return {
        "message": "Profile updated and recommendations generated",
        "recommendations": recommendations
    }






















# Create or Update Agent
@app.post("/create_or_update_agent/")
def create_or_update_agent(agent: AgentModel, db: Session = Depends(get_db)):
    db_agent = db.query(Agent).filter(Agent.user_id == agent.user_id).first()
    
    if db_agent:
        # Update existing agent
        db_agent.state = agent.state if agent.state else db_agent.state
        db_agent.last_active = agent.last_active if agent.last_active else db_agent.last_active
        db_agent.next_action = agent.next_action if agent.next_action else db_agent.next_action
        db.commit()
        return {"message": "Agent updated successfully"}
    else:
        # Create new agent
        new_agent = Agent(
            user_id=agent.user_id,
            state= agent.state or "initial",
            last_active=agent.last_active,
            next_action=agent.next_action
        )
        db.add(new_agent)
        db.commit()
        return {"message": "Agent created successfully"}

@app.get("/check_scheduled_jobs/")
def check_scheduled_jobs():
    jobs = scheduler.get_jobs()
    return [
        {
            "id": job.id,
            "trigger": str(job.trigger),  # Convert trigger to a string representation
            "next_run": job.next_run_time.strftime("%Y-%m-%d %H:%M:%S") if job.next_run_time else "N/A"
        }
        for job in jobs
    ]

# Upload Document
@app.post("/upload_document/")
async def upload_document(
    user_id: int = Form(...), 
    file: UploadFile = File(...), 
    db: Session = Depends(get_db)
):
    try:
        print(f"Received user_id: {user_id}")
        print(f"Received file: {file.filename}")

        # Save the document and get the file path
        file_path = utils.save_document(file, user_id)
        if not file_path:
            raise HTTPException(status_code=500, detail="Error saving document")

        # Store document information in the database
        db_document = models.Document(user_id=user_id, file_path=file_path)
        db.add(db_document)
        db.commit()
        
        return {"message": "Document uploaded successfully", "file_path": file_path}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading document: {str(e)}")



# Improved endpoint for agent auto-training

@app.post("/auto_train_agent/")
async def auto_train_endpoint(
    background_tasks: BackgroundTasks,
    user_id: int = Form(...),
    github : str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    # Fetch or create the agent for the given user_id
    agent = db.query(models.Agent).filter(models.Agent.user_id == user_id).first()
    
    if not agent:
        # If agent doesn't exist, create a new one
        agent = models.Agent(
            user_id=user_id,
            state="initial",
            last_active=None,
            next_action=None
        )
        db.add(agent)
        db.commit()
    
    try:
        # Save the document file dynamically
        file_path = utils.save_document(file, user_id)
        if not file_path:
            raise HTTPException(status_code=500, detail="Error saving document.")
        
        # Schedule auto-training in the background
        background_tasks.add_task(auto_train_agent, db, agent, github, file_path)

        return {"message": "Auto-training scheduled successfully."}
    
    except Exception as e:
        logger.error(f"Error in auto-train endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail="Error scheduling auto-training.")

# Get User Information
@app.get("/user/{user_id}", response_model=UserResponse)
def get_user(user_id: int, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")
    return user


class SubscriptionCreate(BaseModel):
    user_id: int
    payment_id: str
    plan_type: str

@app.post("/api/store_subscription")
async def store_subscription(subscription: SubscriptionCreate, db: Session = Depends(get_db)):
    new_subscription = Subscription(
        user_id=subscription.user_id,
        payment_id=subscription.payment_id,
        plan_type=subscription.plan_type
    )
    db.add(new_subscription)
    try:
        db.commit()
        db.refresh(new_subscription)
        return {"message": "Subscription recorded successfully!", "subscription": new_subscription}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()






# Store all active connections
connected_clients = set()

async def handle_connection(websocket, path):
    # Add the new connection to the set
    connected_clients.add(websocket)
    
    try:
        # Keep listening for messages from the client
        async for message in websocket:
            data = json.loads(message)
            message_type = data.get('type')
            payload = data.get('payload')
            recipient = data.get('to')  # Define recipient logic if needed
            
            # Broadcast the message to all other connected clients
            for client in connected_clients:
                if client != websocket and client.open:
                    await client.send(json.dumps({
                        "type": message_type,
                        "payload": payload
                    }))
    except websockets.exceptions.ConnectionClosed as e:
        print(f"Connection closed: {e}")
    finally:
        # Remove the connection when it's closed
        connected_clients.remove(websocket)

# Start the WebSocket server
async def start_server():
    async with websockets.serve(handle_connection, "localhost", 8000):
        print("WebSocket server started on ws://localhost:8000")
        await asyncio.Future()  # Run forever







# Retrain model periodically on startup
@app.on_event("startup")
def startup_event():
    scheduler.start()
    init_db()
    db = next(get_db())
    retrain_model(db)
    
@sio.event
async def message(sid, data):
    print(f"Message from {sid}: {data}")
    await sio.send(sid, "Message received")



@app.on_event("shutdown")
def on_shutdown():
    scheduler.shutdown()

   

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(sio_app, host="0.0.0.0", port=8000)
    # asyncio.run(start_server())      
